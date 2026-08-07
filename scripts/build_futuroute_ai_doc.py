from __future__ import annotations

import re
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "FutuRoute-AI-Sistem-Mimarisi.md"
OUTPUT = ROOT / "docs" / "FutuRoute-Yapay-Zeka-Calisma-Prensipleri.docx"

INDIGO = "4F46E5"
PURPLE = "7C3AED"
NAVY = "111827"
SLATE = "334155"
MUTED = "64748B"
LIGHT = "EEF2FF"
PALE = "F8FAFC"
WHITE = "FFFFFF"
GREEN = "059669"
AMBER = "D97706"


def set_cell_fill(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_text(cell, text: str, *, bold=False, color=NAVY, size=8.5) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Aptos"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)


def add_field(run, instruction: str) -> None:
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instruction
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr_text, fld_char2])


def add_hyperlink(paragraph, text: str, url: str):
    part = paragraph.part
    relationship_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    new_run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), INDIGO)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_properties.extend([color, underline])
    new_run.append(run_properties)
    text_node = OxmlElement("w:t")
    text_node.text = text
    new_run.append(text_node)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


INLINE_PATTERN = re.compile(r"(\*\*.+?\*\*|`.+?`|\[[^\]]+\]\([^)]+\))")


def add_inline(paragraph, text: str) -> None:
    text = text.replace("  ", " ").strip()
    position = 0
    for match in INLINE_PATTERN.finditer(text):
        if match.start() > position:
            paragraph.add_run(text[position:match.start()])
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
            run.font.color.rgb = RGBColor.from_string(NAVY)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Cascadia Mono"
            run.font.size = Pt(8.2)
            run.font.color.rgb = RGBColor.from_string(PURPLE)
        else:
            link = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token)
            if link:
                add_hyperlink(paragraph, link.group(1), link.group(2))
        position = match.end()
    if position < len(text):
        paragraph.add_run(text[position:])


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(1.9)
    section.right_margin = Cm(1.9)
    section.header_distance = Cm(0.7)
    section.footer_distance = Cm(0.7)

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(9.3)
    normal.font.color.rgb = RGBColor.from_string(SLATE)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08

    for style_name, size, color, before, after in (
        ("Title", 30, NAVY, 0, 10),
        ("Heading 1", 20, NAVY, 16, 7),
        ("Heading 2", 14, INDIGO, 12, 5),
        ("Heading 3", 11, PURPLE, 9, 3),
    ):
        style = doc.styles[style_name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name, base, color, fill in (
        ("Executive Quote", "Normal", NAVY, LIGHT),
        ("Source Note", "Normal", MUTED, PALE),
        ("Caption Custom", "Normal", MUTED, WHITE),
    ):
        if name not in doc.styles:
            style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
            style.base_style = doc.styles[base]
        style = doc.styles[name]
        style.font.name = "Aptos"
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_after = Pt(6)
        if name == "Executive Quote":
            style.font.size = Pt(12)
            style.font.italic = True
            style.paragraph_format.left_indent = Cm(0.45)
            style.paragraph_format.right_indent = Cm(0.45)
        elif name == "Caption Custom":
            style.font.size = Pt(8)
            style.font.italic = True


def add_header_footer(doc: Document) -> None:
    for section in doc.sections:
        header = section.header
        p = header.paragraphs[0]
        p.text = "FutuRoute  /  AI ÇALIŞMA PRENSİPLERİ"
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.runs[0].font.name = "Aptos"
        p.runs[0].font.size = Pt(7.5)
        p.runs[0].font.bold = True
        p.runs[0].font.color.rgb = RGBColor.from_string(INDIGO)

        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("FutuRoute · Gizlilik kontrollü sunum destek dokümanı   |   ")
        run.font.name = "Aptos"
        run.font.size = Pt(7.5)
        run.font.color.rgb = RGBColor.from_string(MUTED)
        add_field(p.add_run(), "PAGE")


def add_cover(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(0.38)
    table.columns[1].width = Cm(16.2)
    set_cell_fill(table.cell(0, 0), INDIGO)
    set_cell_fill(table.cell(0, 1), WHITE)
    left = table.cell(0, 1)
    left.text = ""
    p = left.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    r = p.add_run("FutuRoute")
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor.from_string(INDIGO)
    p = left.add_paragraph()
    p.paragraph_format.space_before = Pt(48)
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run("Yapay Zekâ\nÇalışma Prensipleri")
    r.bold = True
    r.font.name = "Aptos Display"
    r.font.size = Pt(30)
    r.font.color.rgb = RGBColor.from_string(NAVY)
    p = left.add_paragraph()
    r = p.add_run("Güvenilirlik · Öğrenci güvenliği · Doğrulanmış katalog · İnsan gözetimi")
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor.from_string(PURPLE)

    p = left.add_paragraph()
    p.paragraph_format.space_before = Pt(28)
    p.paragraph_format.left_indent = Cm(0.4)
    p.paragraph_format.right_indent = Cm(0.7)
    r = p.add_run("“Yapay zekâ öğrenci adına karar vermez; belirsiz bir isteği güvenli ve düzenlenebilir bir sonraki adıma dönüştürür.”")
    r.italic = True
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor.from_string(SLATE)

    chips = left.add_table(rows=1, cols=3)
    chips.autofit = False
    for index, (label, color) in enumerate((("3 DAR GÖREV", INDIGO), ("0 HAM PROMPT LOGU", PURPLE), ("İNSAN ONAYI", GREEN))):
        cell = chips.cell(0, index)
        set_cell_fill(cell, color)
        set_cell_text(cell, label, bold=True, color=WHITE, size=8)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = left.add_paragraph()
    p.paragraph_format.space_before = Pt(84)
    r = p.add_run("Yönetici ve yatırımcı sunumları için düzenlenebilir destek dokümanı")
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string(MUTED)
    p = left.add_paragraph()
    r = p.add_run("Sürüm 1.0  ·  7 Ağustos 2026")
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string(NAVY)
    doc.add_page_break()


def add_contents(doc: Document, headings: Iterable[str]) -> None:
    doc.add_heading("İçindekiler", level=1)
    p = doc.add_paragraph("Sunum sırasında bölümler bağımsız konuşma notu olarak da kullanılabilir.")
    p.style = "Source Note"
    for number, heading in enumerate(headings, 1):
        table = doc.add_table(rows=1, cols=2)
        table.autofit = False
        table.columns[0].width = Cm(1.1)
        table.columns[1].width = Cm(14.9)
        set_cell_fill(table.cell(0, 0), INDIGO if number % 2 else PURPLE)
        set_cell_text(table.cell(0, 0), f"{number:02d}", bold=True, color=WHITE, size=9)
        table.cell(0, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_fill(table.cell(0, 1), PALE)
        set_cell_text(table.cell(0, 1), heading, bold=True, color=NAVY, size=9)
        doc.add_paragraph().paragraph_format.space_after = Pt(0)
    doc.add_page_break()


def add_architecture_diagram(doc: Document) -> None:
    stages = [
        ("01", "Öğrenci girdisi", "Açık hedef metni · öğrenci kontrolü", INDIGO),
        ("02", "Sunucu güven katmanı", "STUDENT rolü · sahiplik · Zod · rate limit", PURPLE),
        ("03", "Veri minimizasyonu", "Sınıf aralığı · kontrollü etiketler · RIASEC", INDIGO),
        ("04", "Görev tanımlı AI geçidi", "Qwen3.5-9B · non-thinking · token bütçesi", PURPLE),
        ("05", "Çıktı doğrulama", "JSON Schema + Zod · güvenlik politikası", INDIGO),
        ("06", "Kaynak modu", "AI önerisi / Hazır şablon / Kural tabanlı", PURPLE),
        ("07", "İnsan onayı", "Düzenleme · onay · rehber/veli desteği", GREEN),
        ("08", "Kalıcı kayıt ve ölçüm", "Ham metinsiz telemetri · geri bildirim", SLATE),
    ]
    diagram = doc.add_table(rows=len(stages), cols=3)
    diagram.alignment = WD_TABLE_ALIGNMENT.CENTER
    diagram.autofit = False
    diagram.columns[0].width = Cm(1.1)
    diagram.columns[1].width = Cm(5.4)
    diagram.columns[2].width = Cm(9.2)
    for row, (number, title, detail, color) in zip(diagram.rows, stages):
        set_cell_fill(row.cells[0], color)
        set_cell_text(row.cells[0], number, bold=True, color=WHITE, size=9)
        row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_fill(row.cells[1], LIGHT)
        set_cell_text(row.cells[1], title, bold=True, color=NAVY, size=9)
        set_cell_fill(row.cells[2], PALE)
        set_cell_text(row.cells[2], detail, color=SLATE, size=8.5)
    caption = doc.add_paragraph("Şekil 1. Normal akış. Her aşama modelden bağımsız bir kontrol noktasıdır; hata durumunda güvenli kaynak modu kullanıcıya açıkça gösterilir.")
    caption.style = "Caption Custom"
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_markdown_table(doc: Document, lines: list[str]) -> None:
    rows = []
    for index, line in enumerate(lines):
        cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
        if index == 1 and all(set(cell) <= {"-", ":"} for cell in cells):
            continue
        rows.append(cells)
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_repeat_table_header(table.rows[0])
    for row_index, values in enumerate(rows):
        for col_index, value in enumerate(values):
            cell = table.cell(row_index, col_index)
            if row_index == 0:
                set_cell_fill(cell, NAVY)
                set_cell_text(cell, value.replace("`", ""), bold=True, color=WHITE, size=8.2)
            else:
                set_cell_fill(cell, WHITE if row_index % 2 else PALE)
                set_cell_text(cell, value.replace("`", ""), color=SLATE, size=8.1)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


PAGE_BREAK_HEADINGS = {
    "Uçtan uca çalışma akışı",
    "Tek sayfalık mimari görünümü",
    "RIASEC ve kişilik envanterlerinin rolü",
    "Mahremiyet, çocuk güvenliği ve insan gözetimi",
    "Kalite ölçütleri ve pilot yayın kapıları",
    "Sık sorulan sorular",
    "Sunum için 60 saniyelik konuşma metni",
    "Kaynaklar",
}


def parse_body(doc: Document, lines: list[str]) -> None:
    index = 0
    in_mermaid = False
    while index < len(lines):
        raw = lines[index].rstrip()
        line = raw.strip()
        if not line or line == "---":
            index += 1
            continue
        if line.startswith("```mermaid"):
            in_mermaid = True
            add_architecture_diagram(doc)
            index += 1
            continue
        if in_mermaid:
            if line.startswith("```"):
                in_mermaid = False
            index += 1
            continue
        if line.startswith("|"):
            table_lines = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index].strip())
                index += 1
            add_markdown_table(doc, table_lines)
            continue
        if line.startswith("## "):
            title = line[3:].strip()
            if title in PAGE_BREAK_HEADINGS:
                doc.add_page_break()
            doc.add_heading(title, level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
        elif line.startswith("# "):
            pass
        elif line.startswith("> "):
            p = doc.add_paragraph(style="Executive Quote")
            add_inline(p, line[2:])
            p.paragraph_format.keep_together = True
        elif match := re.match(r"^(\d+)\. (.*)", line):
            # Preserve the numbering written in Markdown. Word's built-in list
            # style continues numbering across sections, which is undesirable
            # for independently numbered presentation lists.
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.55)
            p.paragraph_format.first_line_indent = Cm(-0.55)
            number = p.add_run(f"{match.group(1)}. ")
            number.bold = True
            add_inline(p, match.group(2))
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, line[2:])
        else:
            paragraph_lines = [line]
            while index + 1 < len(lines):
                next_line = lines[index + 1].strip()
                if not next_line or next_line.startswith(("#", ">", "- ", "|", "```")) or re.match(r"^\d+\. ", next_line):
                    break
                paragraph_lines.append(next_line)
                index += 1
            p = doc.add_paragraph()
            add_inline(p, " ".join(paragraph_lines))
            if line.startswith("**Durum notu:"):
                p.style = "Source Note"
        index += 1


def set_document_metadata(doc: Document) -> None:
    props = doc.core_properties
    props.title = "FutuRoute Yapay Zekâ Çalışma Prensipleri"
    props.subject = "Güvenilirlik, öğrenci güvenliği ve sistem mimarisi"
    props.author = "FutuRoute"
    props.last_modified_by = "FutuRoute"
    props.keywords = "FutuRoute, yapay zekâ, öğrenci güvenliği, Qwen, RIASEC"
    props.comments = ""


def main() -> None:
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    body_start = next(index for index, line in enumerate(lines) if line.startswith("## Bir dakikalık"))
    body_lines = lines[body_start:]
    headings = [line[3:].strip() for line in body_lines if line.startswith("## ")]

    doc = Document()
    configure_document(doc)
    add_cover(doc)
    add_contents(doc, headings)
    parse_body(doc, body_lines)
    add_header_footer(doc)
    set_document_metadata(doc)

    # Accessibility: document language and compatibility mode.
    settings = doc.settings._element
    theme_lang = OxmlElement("w:themeFontLang")
    theme_lang.set(qn("w:val"), "tr-TR")
    settings.append(theme_lang)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
